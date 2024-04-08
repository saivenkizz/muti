import json
import os
import numpy as np
import streamlit as st
from streamlit import session_state
import pdfplumber
import docx
import pyaes
import random
import pandas as pd
import base64
import hashlib

session_state = st.session_state
if "user_index" not in st.session_state:
    st.session_state["user_index"] = 0


def generateKey(user_key, admin_auth, token_auth):
    key = hashlib.sha256(user_key.encode("utf-8") + admin_auth.encode("utf-8") + token_auth.encode("utf-8")).digest()[:16]
    return key

def signup(json_file_path="data.json"):
    st.title("Data Publisher signup Page")
    with st.form("signup_form"):
        st.write("Fill in the details below to create an account:")
        name = st.text_input("Name:")
        email = st.text_input("Email:")
        age = st.number_input("Age:", min_value=0, max_value=120)
        sex = st.radio("Sex:", ("Male", "Female", "Other"))
        password = st.text_input("Password:")
        confirm_password = st.text_input("Confirm Password:")

        if st.form_submit_button("Signup"):
            if password == confirm_password:
                user = create_account(name, email, age, sex, password, json_file_path)
                session_state["logged_in"] = True
                session_state["user_info"] = user
            else:
                st.error("Passwords do not match. Please try again.")


def check_login(username, password, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)

        for user in data["users"]:
            if user["email"] == username and user["password"] == password:
                session_state["logged_in"] = True
                session_state["user_info"] = user
                st.success("Login successful!")
                return user

        st.error("Invalid credentials. Please try again.")
        return None
    except Exception as e:
        st.error(f"Error checking login: {e}")
        return None


def initialize_database():
    try:
        if not os.path.exists("data.json"):
            data = {"users": []}
            with open("data.json", "w") as json_file:
                json.dump(data, json_file)
                
    except Exception as e:
        print(f"Error initializing database: {e}")


def create_account(name, email, age, sex, password, json_file_path="data.json"):
    try:
        if not os.path.exists(json_file_path) or os.stat(json_file_path).st_size == 0:
            data = {"users": []}
        else:
            with open(json_file_path, "r") as json_file:
                data = json.load(json_file)

        # Append new user data to the JSON structure
        user_info = {
            "name": name,
            "email": email,
            "age": age,
            "sex": sex,
            "password": password,
            "files":None,
        }
        data["users"].append(user_info)

        with open(json_file_path, "w") as json_file:
            json.dump(data, json_file, indent=4)

        st.success("Account created successfully! You can now login.")
        return user_info
    except json.JSONDecodeError as e:
        st.error(f"Error decoding JSON: {e}")
        return None
    except Exception as e:
        st.error(f"Error creating account: {e}")
        return None


def login(json_file_path="data.json"):
    st.title("Login Page")
    username = st.text_input("Username:")
    password = st.text_input("Password:")

    login_button = st.button("Login")

    if login_button:
        user = check_login(username, password, json_file_path)
        if user is not None:
            session_state["logged_in"] = True
            session_state["user_info"] = user
        else:
            st.error("Invalid credentials. Please try again.")


def get_user_info(email, json_file_path="data.json"):
    try:
        with open(json_file_path, "r") as json_file:
            data = json.load(json_file)
            for user in data["users"]:
                if user["email"] == email:
                    return user
        return None
    except Exception as e:
        st.error(f"Error getting user information: {e}")
        return None
def extract_text(file) -> str:
    if isinstance(file, str):
        file_extension = os.path.splitext(file)[1].lower()
    else:
        file_extension = os.path.splitext(file.name)[1].lower()
    
    if file_extension == '.pdf':
        if isinstance(file, str):
            with pdfplumber.open(file) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
        else:
            with pdfplumber.open(file) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_extension == '.docx':
        if isinstance(file, str):
            doc = docx.Document(file)
        else:
            doc = docx.Document(file)
        text = '\n'.join([para.text for para in doc.paragraphs])
    else:
        if isinstance(file, str):
            with open(file, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            with file as f:
                text = f.read()
    return text

def render_dashboard(user_info, json_file_path="data.json"):
    try:
        st.title(f"Welcome to the Dashboard, {user_info['name']}!")
        st.subheader("Data Publisher Information:")
        st.write(f"Name: {user_info['name']}")
        st.write(f"Sex: {user_info['sex']}")
        st.write(f"Age: {user_info['age']}")
        st.image("data_publisher.jpg")
    except Exception as e:
        st.error(f"Error rendering dashboard: {e}")

def get_keys():
    with st.form("credentials"):
        st.write("Enter the credentials to encrypt the file:")
        user_key = st.text_input("Enter the User key:")
        admin_auth = st.text_input("Enter the Admin Authentication key:")
        token_auth = st.text_input("Enter the Token Authentication key:")
        if st.form_submit_button("Encrypt and Upload"):
            return user_key, admin_auth, token_auth
    return None, None, None
def main(
    json_file_path="data.json",
):
    st.sidebar.title("Secure Multi-Factor Authentication")
    page = st.sidebar.radio(
        "Go to",
        (
            "Signup/Login",
            "Dashboard",
            "File Upload",
            "File Download",
        ),
        key="Secure Multi-Factor Authentication",
    )

    if page == "Signup/Login":
        st.title("Signup/Login Page")
        login_or_signup = st.radio(
            "Select an option", ("Login", "Signup"), key="login_signup"
        )
        if login_or_signup == "Login":
            login(json_file_path)
        else:
            signup(json_file_path)

    elif page == "Dashboard":
        if session_state.get("logged_in"):
            render_dashboard(session_state["user_info"])
        else:
            st.warning("Please login/signup to view the dashboard.")

    elif page == "File Upload":
        if session_state.get("logged_in"):
            st.title("File Upload")
            uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "png", "jpg", "jpeg"])
            if uploaded_file is not None:
                file_details = {
                    "filename": uploaded_file.name,
                    "filetype": uploaded_file.type,
                    "filesize": uploaded_file.size,
                }
                st.write("Name: %s" % uploaded_file.name)
                st.write("Type: %s" % uploaded_file.type)
                st.write("Size: %s" % uploaded_file.size)
                st.write("Enter the credentials to encrypt the file:")
                with st.form("credentials"):
                    st.write("Enter the credentials to encrypt the file:")
                    user_key = st.text_input("Enter the User key:")
                    admin_auth = st.text_input("Enter the Admin Authentication key:")
                    token_auth = st.text_input("Enter the Token Authentication key:")
                    if st.form_submit_button("Encrypt and Upload"):
                        with open(json_file_path, "r+") as json_file:
                            print("1")
                            file = base64.b64encode(uploaded_file.read()).decode("utf-8")
                            data = json.load(json_file)
                            user_index = next((i for i, user in enumerate(data["users"]) if user["email"] ==session_state["user_info"]["email"]), None)
                            if user_index is not None:
                                user_info = data["users"][user_index]
                                if user_info["files"] is None:
                                    user_info["files"] = []
                                key = generateKey(user_key, admin_auth, token_auth)
                                aes = pyaes.AESModeOfOperationCTR(key)  
                                cipher_text = aes.encrypt(file)
                                cipher_text = base64.b64encode(cipher_text).decode("utf-8")
                                current_time = str(np.datetime64('now'))
                                file_name = uploaded_file.name
                                for file in user_info["files"]:
                                    if file["file"] == file_name:
                                        file_name = file_name.split(".")[0] + "_1." + file_name.split(".")[1]
                                user_info["files"].append({"file": uploaded_file.name, "data": cipher_text, "time": current_time, "sanitized": False})
                                session_state["user_info"] = user_info
                                json_file.seek(0)
                                json.dump(data, json_file, indent=4)
                                json_file.truncate()
                        
                        st.success("File uploaded successfully!")
        else:
            st.warning("Please login/signup to access this page.")
            
            
    elif page == "File Download":
        if session_state.get("logged_in"):
            st.title("File Download")   
            i = 1
            user_info = session_state["user_info"]
            if len(session_state["user_info"]["files"]) == 0:
                st.warning("No files uploaded yet.")
                return
            for file in session_state["user_info"]["files"]:
                files = []
                file_data = {}
                file_data["S.No"] = i
                file_data["File Name"] = file["file"]
                file_data["Upload Time"] = file["time"]
                files.append(file_data)
                i += 1
                st.table(files)
                try:
                    with st.form("credentials2"):
                        st.write("Enter the credentials to decrypt the file:")
                        user_key = st.text_input("Enter the User key:")
                        admin_auth = st.text_input("Enter the Admin Authentication key:")
                        token_auth = st.text_input("Enter the Token Authentication key:")
                        if st.form_submit_button("Decrypt and Download"):
                            key = generateKey(user_key, admin_auth, token_auth)
                            aes = pyaes.AESModeOfOperationCTR(key)
                            data = base64.b64decode(file["data"])
                            decrypted_text = aes.decrypt(data).decode('utf-8')
                            data = base64.b64decode(decrypted_text)
                            with open(file["file"], "wb") as f:
                                f.write(data)
                                st.success("File downloaded successfully!")
                except Exception as e:
                    st.error(f"Wrong credentials")
        else:
            st.warning("Please login/signup to access this page.")
    


if __name__ == "__main__":

    initialize_database()
    main()
